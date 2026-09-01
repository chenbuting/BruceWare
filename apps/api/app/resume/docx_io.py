"""按充实版简历样式导出 / 解析 Word。"""

from __future__ import annotations

import io
import json
import re
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips

from app.resume.text import flatten_resume

BLUE = "5B9BD5"
NAME_COLOR = "1F4E79"
TITLE_COLOR = "262626"
FONT = "微软雅黑"
SECTIONS = ("个人优势", "自我评价", "工作经历", "个人技能", "项目经验", "项目经历")
JOB_HEAD = re.compile(
    r"^(20\d{2}[./年]\d{1,2}(?:[./月]\d{1,2})?\s*[-~—至到]{1,2}\s*(?:今|(?:20\d{2}[./年]\d{1,2}(?:[./月]\d{1,2})?))?)(.*)$"
)
COMPANY_TAIL = re.compile(r"^(.+?(?:有限责任公司|有限公司|股份公司|集团|公司))(.*)$")


def _set_run(run, size: float, bold: bool = False, color: str | None = None) -> None:
    run.font.name = FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _spacing(paragraph, after: int = 60) -> None:
    """行距和段后，单位 twips，对齐充实版。"""

    pPr = paragraph._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:spacing")):
        pPr.remove(old)
    el = OxmlElement("w:spacing")
    el.set(qn("w:before"), "0")
    el.set(qn("w:after"), str(after))
    el.set(qn("w:line"), "269")
    el.set(qn("w:lineRule"), "auto")
    pPr.append(el)


def _shade(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _no_borders(table) -> None:
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), "nil")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)


def _set_cell_width(cell, twips: int) -> None:
    cell.width = Twips(twips)
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcW")):
        tcPr.remove(old)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(twips))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _add_text(paragraph, text: str, size: float, bold: bool = False, color: str | None = None, after: int = 60) -> None:
    run = paragraph.add_run(text)
    _set_run(run, size, bold, color)
    _spacing(paragraph, after)


def _section_bar(doc: Document, title: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    _no_borders(table)
    left, right = table.rows[0].cells
    _set_cell_width(left, 1560)
    _set_cell_width(right, 8580)
    _shade(left, BLUE)
    left.text = ""
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_text(p, title, 11, True, "FFFFFF", 0)
    right.text = ""


def _pair(label_a: str, value_a: str, label_b: str, value_b: str) -> str:
    return f"{label_a}：{value_a or '　'}     {label_b}：{value_b or '　'}"


def _load_form(raw: str) -> dict[str, Any]:
    text = (raw or "").strip()
    if not text:
        return {"version": 1, "basic": {}, "summary": "", "jobs": [], "skills": [], "projects": []}
    try:
        data = json.loads(text)
        if isinstance(data, dict) and data.get("version") == 1:
            return data
    except Exception:
        pass
    return {"version": 1, "basic": {}, "summary": flatten_resume(text) or text, "jobs": [], "skills": [], "projects": []}


def build_docx(title: str, raw: str) -> bytes:
    """按充实版版式生成 Word。"""

    data = _load_form(raw)
    basic = data.get("basic") if isinstance(data.get("basic"), dict) else {}
    doc = Document()
    section = doc.sections[0]
    section.page_width = Twips(11900)
    section.page_height = Twips(16840)
    section.top_margin = Twips(720)
    section.bottom_margin = Twips(720)
    section.left_margin = Twips(720)
    section.right_margin = Twips(720)

    head = doc.add_table(rows=1, cols=2)
    head.autofit = False
    _no_borders(head)
    info_cell, photo_cell = head.rows[0].cells
    _set_cell_width(info_cell, 7083)
    _set_cell_width(photo_cell, 3059)
    info_cell.text = ""
    photo_cell.text = ""

    name = str(basic.get("name") or title or "简历").strip()
    p_name = info_cell.paragraphs[0]
    _add_text(p_name, name, 24, True, NAME_COLOR, 80)

    school = str(basic.get("school") or "").strip()
    if str(basic.get("grad_year") or "").strip():
        year = str(basic.get("grad_year")).strip()
        school = f"{school}（{year}年毕业）" if school else f"{year}年毕业"
    lines = [
        _pair("毕业院校", school, "学    历", str(basic.get("education") or "")),
        _pair("年    龄", str(basic.get("age") or ""), "专    业", str(basic.get("major") or "")),
        _pair("性    别", str(basic.get("gender") or ""), "联系电话", str(basic.get("phone") or "")),
        f"邮    箱：{str(basic.get('email') or '')}",
    ]
    if str(basic.get("hometown") or "").strip():
        lines.append(f"籍    贯：{str(basic.get('hometown')).strip()}")
    if str(basic.get("target_job") or "").strip():
        lines.append(f"求职意向：{str(basic.get('target_job')).strip()}")
    for line in lines:
        p = info_cell.add_paragraph()
        _add_text(p, line, 10.5, False, after=40)

    _section_bar(doc, "个人优势")
    summary = str(data.get("summary") or "").strip()
    if summary:
        blocks = [item.strip() for item in summary.split("\n") if item.strip()]
        for index, block in enumerate(blocks):
            p = doc.add_paragraph()
            _add_text(p, block, 10.5, True, after=160 if index == len(blocks) - 1 else 80)

    jobs = data.get("jobs") if isinstance(data.get("jobs"), list) else []
    if jobs:
        _section_bar(doc, "工作经历")
        for job in jobs:
            if not isinstance(job, dict):
                continue
            head_line = f"{str(job.get('period') or '').strip()}{str(job.get('company') or '').strip()}{str(job.get('role') or '').strip()}"
            if head_line:
                p = doc.add_paragraph()
                _add_text(p, head_line, 12, True, TITLE_COLOR, 100)
            intro = str(job.get("intro") or "").strip()
            if intro:
                p = doc.add_paragraph()
                _add_text(p, intro, 10.5, True, after=60)
            bullets = [str(item).strip() for item in (job.get("bullets") or []) if str(item).strip()]
            for index, item in enumerate(bullets, start=1):
                p = doc.add_paragraph()
                _add_text(p, f"{index}、{item}", 10.5, True, after=60)

    skills = data.get("skills") if isinstance(data.get("skills"), list) else []
    if skills:
        _section_bar(doc, "个人技能")
        for skill in skills:
            if not isinstance(skill, dict):
                continue
            title_line = str(skill.get("title") or "").strip()
            if title_line:
                p = doc.add_paragraph()
                _add_text(p, title_line, 10.5, True, after=60)
            bullets = [str(item).strip() for item in (skill.get("bullets") or []) if str(item).strip()]
            for index, item in enumerate(bullets, start=1):
                p = doc.add_paragraph()
                _add_text(p, f"{index}、{item}", 10.5, True, after=60)

    projects = data.get("projects") if isinstance(data.get("projects"), list) else []
    if projects:
        _section_bar(doc, "项目经验")
        for project in projects:
            if not isinstance(project, dict):
                continue
            name = str(project.get("name") or "").strip()
            if name:
                p = doc.add_paragraph()
                _add_text(p, name, 10.5, True, after=60)
            stack = str(project.get("stack") or "").strip()
            if stack:
                p = doc.add_paragraph()
                _add_text(p, f"技术栈：{stack}", 10.5, True, after=60)
            goal = str(project.get("goal") or "").strip()
            if goal:
                p = doc.add_paragraph()
                _add_text(p, f"项目描述：{goal}", 10.5, True, after=60)
            duties = [str(item).strip() for item in (project.get("duties") or []) if str(item).strip()]
            if duties:
                p = doc.add_paragraph()
                _add_text(p, "个人职责：", 10.5, True, after=60)
                for index, item in enumerate(duties, start=1):
                    p = doc.add_paragraph()
                    _add_text(p, f"{index}、{item}", 10.5, True, after=60)
            result = str(project.get("result") or "").strip()
            if result:
                p = doc.add_paragraph()
                _add_text(p, f"项目成果：{result}", 10.5, True, after=60)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _xml_text(el) -> str:
    lines: list[str] = []
    for para in el.iter(qn("w:p")):
        line = "".join(node.text or "" for node in para.iter(qn("w:t"))).strip()
        if line:
            lines.append(line)
    if lines:
        return "\n".join(lines)
    return "".join(node.text or "" for node in el.iter(qn("w:t"))).strip()


def parse_docx(data: bytes) -> dict[str, Any]:
    """把充实版一类 Word 拆回栏目。"""

    doc = Document(io.BytesIO(data))
    blocks: list[tuple[str, str]] = []
    for child in doc.element.body:
        tag = child.tag.split("}")[-1]
        if tag == "tbl":
            header_parts: list[str] = []
            section = ""
            for row in child.findall(qn("w:tr")):
                for cell in row.findall(qn("w:tc")):
                    text = _xml_text(cell)
                    if not text:
                        continue
                    if text in SECTIONS:
                        section = text
                    else:
                        header_parts.append(text)
            if header_parts:
                blocks.append(("header", "\n".join(header_parts)))
            if section:
                blocks.append(("section", section))
        elif tag == "p":
            text = _xml_text(child)
            if text:
                blocks.append(("p", text))

    form: dict[str, Any] = {
        "version": 1,
        "basic": {
            "name": "",
            "school": "",
            "grad_year": "",
            "education": "",
            "age": "",
            "major": "",
            "hometown": "",
            "gender": "",
            "phone": "",
            "email": "",
            "target_job": "",
        },
        "summary": "",
        "jobs": [],
        "skills": [],
        "projects": [],
    }
    current = ""
    bucket: list[str] = []

    def flush() -> None:
        nonlocal bucket, current
        if current in ("个人优势", "自我评价"):
            form["summary"] = "\n".join(bucket).strip()
        elif current == "工作经历":
            form["jobs"] = _parse_jobs(bucket)
        elif current == "个人技能":
            form["skills"] = _parse_skills(bucket)
        elif current in ("项目经验", "项目经历"):
            form["projects"] = _parse_projects(bucket)
        bucket = []

    for kind, text in blocks:
        if kind == "header":
            _fill_basic(form["basic"], text)
            continue
        if kind == "section":
            flush()
            current = text
            continue
        if not current and kind == "p":
            _fill_basic(form["basic"], text)
            if not form["basic"]["name"] and "：" not in text and ":" not in text:
                form["basic"]["name"] = text[:20]
            continue
        bucket.append(text)
    flush()

    if not form["basic"]["name"]:
        form["basic"]["name"] = "导入的简历"
    return form


def _pick(text: str, *labels: str) -> str:
    for label in labels:
        match = re.search(label + r"[：:]\s*([^\n]+)", text)
        if match:
            value = re.split(r"学\s*历|专\s*业|年\s*龄|性\s*别|籍\s*贯|联系电话|电\s*话|邮\s*箱|求职意向", match.group(1), maxsplit=1)[0]
            return value.strip().strip("　 ")
    return ""


def _fill_basic(basic: dict[str, str], text: str) -> None:
    school = _pick(text, r"毕业院校")
    if school:
        year = re.search(r"（?((?:19|20)\d{2})年?毕业", school)
        if year:
            basic["grad_year"] = year.group(1)
            school = re.sub(r"[（(].*?毕业[)）]", "", school).strip()
        basic["school"] = school
    basic["education"] = basic["education"] or _pick(text, r"学\s*历")
    basic["age"] = basic["age"] or _pick(text, r"年\s*龄")
    basic["major"] = basic["major"] or _pick(text, r"专\s*业")
    basic["gender"] = basic["gender"] or _pick(text, r"性\s*别")
    basic["hometown"] = basic["hometown"] or _pick(text, r"籍\s*贯")
    basic["phone"] = basic["phone"] or _pick(text, r"联系电话", r"电\s*话")
    basic["email"] = basic["email"] or _pick(text, r"邮\s*箱")
    basic["target_job"] = basic["target_job"] or _pick(text, r"求职意向")
    for line in text.splitlines():
        clean = line.strip()
        if clean and "：" not in clean and ":" not in clean and not clean.startswith("毕业") and 1 < len(clean) <= 12:
            basic["name"] = basic["name"] or clean
            break
    if not basic["name"]:
        head = re.split(r"毕业院校|求职意向", re.sub(r"\s+", "", text), maxsplit=1)[0]
        if 1 < len(head) <= 8 and "：" not in head and ":" not in head:
            basic["name"] = head


def _is_bullet(line: str) -> bool:
    return bool(re.match(r"^\d+[\.、．]", line))


def _strip_num(line: str) -> str:
    return re.sub(r"^\d+[\.、．]\s*", "", line).strip()


def _split_job_head(line: str) -> dict[str, Any] | None:
    match = JOB_HEAD.match(line)
    if not match:
        return None
    current = {"period": match.group(1).strip(), "company": "", "role": "", "intro": "", "bullets": []}
    rest = match.group(2).strip()
    company = COMPANY_TAIL.match(rest)
    if company:
        current["company"] = company.group(1).strip()
        current["role"] = company.group(2).strip()
    else:
        current["company"] = rest
    return current


def _parse_jobs(lines: list[str]) -> list[dict[str, Any]]:
    jobs: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None
    for line in lines:
        started = _split_job_head(line)
        if started:
            if current:
                jobs.append(current)
            current = started
            continue
        if current is None:
            continue
        if _is_bullet(line):
            current["bullets"].append(_strip_num(line))
        elif not current["intro"]:
            current["intro"] = line
        else:
            current["intro"] += line
    if current:
        jobs.append(current)
    return jobs


def _parse_skills(lines: list[str]) -> list[dict[str, Any]]:
    skills: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None
    for line in lines:
        if _is_bullet(line):
            if current is None:
                current = {"title": "", "bullets": []}
            current["bullets"].append(_strip_num(line))
            continue
        if current:
            skills.append(current)
        current = {"title": line, "bullets": []}
    if current:
        skills.append(current)
    return skills


def _parse_projects(lines: list[str]) -> list[dict[str, Any]]:
    projects: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None
    mode = "name"

    def start(name: str) -> None:
        nonlocal current, mode
        if current:
            projects.append(current)
        current = {"name": name, "stack": "", "goal": "", "duties": [], "result": ""}
        mode = "name"

    for line in lines:
        if re.match(r"^技术栈[：:]", line):
            if current is None:
                start("")
            current["stack"] = re.sub(r"^技术栈[：:]\s*", "", line)
            mode = "stack"
            continue
        if re.match(r"^项目(?:描述|介绍)[：:]", line):
            if current is None:
                start("")
            current["goal"] = re.sub(r"^项目(?:描述|介绍)[：:]\s*", "", line)
            mode = "goal"
            continue
        if re.match(r"^个人职责[：:]?", line):
            if current is None:
                start("")
            rest = re.sub(r"^个人职责[：:]\s*", "", line)
            mode = "duties"
            if rest and rest != "个人职责":
                current["duties"].append(rest)
            continue
        if re.match(r"^项目成果[：:]", line):
            if current is None:
                start("")
            current["result"] = re.sub(r"^项目成果[：:]\s*", "", line)
            mode = "result"
            continue
        if _is_bullet(line):
            if current is None:
                start("")
            if mode == "result":
                current["result"] += _strip_num(line)
            else:
                mode = "duties"
                current["duties"].append(_strip_num(line))
            continue
        if current is None or mode in ("duties", "result", "name"):
            start(line)
            continue
        if mode == "goal":
            current["goal"] += line
        elif mode == "stack":
            current["stack"] += line
        else:
            start(line)
    if current:
        projects.append(current)
    return projects
